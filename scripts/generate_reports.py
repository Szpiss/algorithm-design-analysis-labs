from copy import deepcopy
from pathlib import Path

from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT


ROOT = Path("/Users/cuing/algorithm")
TEMPLATE = Path("/Users/cuing/Downloads/学号+姓名 实验一.docx")
REPORT_DIR = ROOT / "reports"


EXPERIMENTS = [
    {
        "filename": "实验一_查找假币问题_实验报告.docx",
        "title": "实验一 查找假币问题",
        "description": (
            "编写一个实验程序用于求解这样的假币问题，共有n(n>3)个硬币，编号为0~n-1，其中有且仅有一个假币，"
            "假币与真币的外观相同但重量不同，不知道假币比真币轻还是重，现在用一架天平称重，天平称重的硬币数没有限制。"
            "最后找出这个假币，使得称重的次数尽可能少。要求采用相关数据进行测试并且输出称重过程。"
        ),
        "question1": "（1）你的解题算法采用的什么算法策略？该算法的主要思想是什么？给出算法框架。",
        "question2": "（2）分析该算法的时间复杂度。",
        "strategy": "分治法。",
        "core": (
            "程序将候选硬币尽量三等分，并通过一次称重得到“左重、右重、平衡”三种互斥结果。"
            "若称重平衡，则假币在未称部分；若不平衡，则可以把问题转化为“某一组中只能是偏重假币，另一组中只能是偏轻假币”的递归子问题。"
            "随着递归进行，候选集合持续缩小，直到确定假币编号和轻重属性。"
        ),
        "framework": [
            "将全部硬币划分为三部分 A、B、C，比较 A 与 B。",
            "若 A 与 B 平衡，则假币一定在 C 中，A 和 B 可视为真币，递归处理 C。",
            "若 A 与 B 不平衡，则把问题转化为“假币可能在一侧偏重，或在另一侧偏轻”的子问题。",
            "对新的候选集合继续三分并递归称量，直到只剩一个候选硬币。",
            "最后借助已经确认的真币判定该硬币是偏重还是偏轻。"
        ],
        "complexity": (
            "每一次有效称重都会把候选状态缩小到原来的约三分之一，因此递归深度约为 log3(n)。"
            "从渐进意义上看，算法时间复杂度可写为 O(log n)，若以天平三种结果来衡量，更准确地说是 O(log3 n)。"
        ),
        "code_summary": (
            "核心程序位于 exp1_fake_coin/src/main.cpp。代码中实现了 searchUnknown、searchPolarized 和 searchKnown 三个递归过程，"
            "分别对应未知状态、偏向已知状态和轻重方向已知状态下的继续求解。"
        ),
        "test_summary": [
            "样例 1：12 枚硬币，7 号假币偏重。程序输出完整称重过程，并最终定位到 7 号硬币。",
            "样例 2：8 枚硬币，2 号假币偏轻。程序可以正确区分出假币编号和偏轻属性。",
            "程序会同时输出实际称重次数，便于和理论下界进行对照分析。"
        ],
        "screenshots": [
            "截图 1：核心递归函数或称重逻辑的代码调试截图。",
            "截图 2：程序运行后展示称重过程的控制台截图。",
            "截图 3：最终定位假币编号与轻重属性的结果截图。"
        ],
        "conclusion": (
            "通过本实验，掌握了分治法中“分解 - 递归求解 - 合并判断”的基本思想，"
            "并体会到如何利用一次比较产生的多种结果来持续缩小搜索空间。"
        ),
    },
    {
        "filename": "实验二_迷宫路径_实验报告.docx",
        "title": "实验二 迷宫路径",
        "description": (
            "编写一个实验程序采用回溯法求解迷宫问题。给定一个m x n个方块的迷宫，每个方块值为0时表示空白，为1时表示障碍物，"
            "在行走时最多只能走到上、下、左、右相邻的方块。求指定入口s到出口t的所有迷宫路径和其中一条最短路径。"
        ),
        "question1": "（1）请给出解空间为子集树的算法框架。",
        "question2": "（2）分析该算法的时间复杂度。",
        "strategy": "回溯法。",
        "core": (
            "程序从入口开始深度优先搜索，每走到一个新位置，就把该位置加入当前路径并标记为已访问；"
            "如果当前位置无法继续扩展或已经到达终点，就回退到上一步，撤销当前选择，继续尝试其他方向。"
            "这样可以枚举所有从入口到出口的可行简单路径，并同步维护其中一条最短路径。"
        ),
        "framework": [
            "从初始状态出发，把当前位置作为当前结点加入路径。",
            "判断当前部分解是否满足约束条件，包括不越界、不进入障碍、不重复访问当前路径中的结点。",
            "若当前结点是终点，则记录一条完整路径，并更新最短路径。",
            "否则按上、下、左、右四个方向扩展子结点，递归搜索。",
            "递归返回时撤销当前位置的访问标记，回溯到父结点。"
        ],
        "complexity": (
            "回溯法需要枚举所有可能路径，在最坏情况下，可行路径数量可能呈指数级增长。"
            "因此该算法的最坏时间复杂度为指数级，常记为 O(4^(m*n)) 的上界形式；"
            "实际运行时会因为障碍物和访问标记的约束而显著减少搜索规模。"
        ),
        "code_summary": (
            "核心程序位于 exp2_maze_paths/src/main.cpp。程序使用 DFS + visited 数组实现回溯，"
            "在到达终点时保存当前路径，并在所有完整路径中维护长度最短的一条。"
        ),
        "test_summary": [
            "测试迷宫共找到 4 条可行路径，程序能够按搜索顺序完整输出每条路径。",
            "程序同时计算最短路径长度，并输出一条对应的最短路径。",
            "结果说明回溯法不仅能求可行解，还能在枚举基础上得到最优解。"
        ],
        "screenshots": [
            "截图 1：DFS 回溯核心代码的调试截图。",
            "截图 2：输出所有路径的运行结果截图。",
            "截图 3：输出最短路径及路径长度的结果截图。"
        ],
        "conclusion": (
            "通过本实验，进一步理解了回溯法的搜索过程、剪枝思想以及恢复现场的重要性，"
            "并掌握了如何在枚举所有解的同时维护最优解。"
        ),
    },
    {
        "filename": "实验三_在原始森林中解救A_实验报告.docx",
        "title": "实验三 在原始森林中解救A",
        "description": (
            "编写一个实验程序解决解救问题，A不幸迷失于原始森林中，B要到原始森林中解救她，他每次只能在上、下、左、右4个方向移动一个单元。"
            "B知道如果遇到金刚他会死的，野狗也会咬他，而且咬了两次（含一只野狗咬两次或两只野狗各咬一次）之后他也会死的。求B能否找到A。"
        ),
        "question1": "（1）请给出队列式分支限界法算法框架。",
        "question2": "（2）结合本章学习，给出优先队列式分支限界法中队列结点类型声明，并解释类型成员表示的含义。",
        "strategy": "分支限界法。",
        "core": (
            "程序把每一个搜索状态定义为“当前位置 + 已被野狗咬的次数”，并使用队列保存所有待扩展的活结点。"
            "每次从队首取出一个活结点，向四个方向扩展；如果越界、遇到金刚、被咬次数达到 2 次，或者该状态已访问过，"
            "则立即剪枝。只要搜索过程中能够到达 A 的位置，就说明 B 可以成功解救 A。"
        ),
        "framework": [
            "初始化活结点队列，将起点状态入队。",
            "当队列非空时，取出队首结点作为当前扩展结点。",
            "若当前结点已经到达目标位置，则输出 Yes 并结束。",
            "按照四个方向生成子结点，并检查约束条件。",
            "对满足条件且未访问过的子结点入队；对不满足条件的子结点直接剪枝。",
            "若队列为空仍未到达目标，则输出 No。"
        ],
        "complexity": (
            "每个位置最多只有两种有效生命状态：未被咬和被咬一次，因此状态总数不超过 2*n*n。"
            "每个状态最多向四个方向扩展一次，所以时间复杂度为 O(n^2)，空间复杂度同样为 O(n^2)。"
        ),
        "code_summary": (
            "核心程序位于 exp3_rescue_a/src/main.cpp。代码使用 queue 存储活结点，"
            "并借助三维 visited[x][y][dogBites] 数组避免重复扩展已经处理过的状态。"
        ),
        "test_summary": [
            "对指导书中的两组样例数据进行测试，程序分别输出 Yes 和 No，与样例答案一致。",
            "程序的剪枝条件包括：越界、进入金刚位置、累计狗咬次数达到 2、以及重复状态。",
            "从运行结果可以看出，该方法能够有效判断是否存在安全可行路径。"
        ],
        "screenshots": [
            "截图 1：队列扩展与剪枝逻辑的代码调试截图。",
            "截图 2：输入指导书样例后的运行结果截图。",
            "截图 3：程序输出 Yes/No 判定结果的截图。"
        ],
        "conclusion": (
            "通过本实验，掌握了分支限界法中“活结点、扩展结点、剪枝约束”的基本概念，"
            "并理解了如何把路径问题转化为带状态约束的图搜索问题。"
        ),
        "extra_question": {
            "title": "优先队列式分支限界法结点类型示例",
            "content": [
                "可以定义结点类型为：struct Node { int x; int y; int dogBites; int depth; int bound; };",
                "其中 x、y 表示当前所在位置坐标；dogBites 表示当前已经被野狗咬的次数；depth 表示从起点走到当前结点所经过的步数；bound 表示该结点的界，可以用估计代价或优先级函数表示，用于决定优先队列中的扩展次序。"
            ],
        },
    },
    {
        "filename": "实验四_赶作业_实验报告.docx",
        "title": "实验四 赶作业",
        "description": (
            "A有n份作业要做，每份作业有一个最后期限，如果在最后期限后交作业老师就会扣分，现在假设完成每份作业都需要一天。"
            "A想安排做作业得顺序，把扣分降到最低，请帮助他实现。"
        ),
        "question1": "（1）通过本实验学习掌握求解调度问题。",
        "question2": "（2）结合本章学习，贪心算法的一般求解过程及算法框架。",
        "strategy": "贪心法。",
        "core": (
            "程序优先考虑“不完成代价最大”的作业。具体做法是先按扣分值从大到小排序，"
            "然后将每份作业安排到不晚于其截止日期的最靠后空闲时间；如果找不到这样的时间槽，就说明这份作业不可避免地会被拖期，"
            "其扣分计入总罚分。为了快速找到最近空闲时间，程序使用并查集维护可用时间槽。"
        ),
        "framework": [
            "计算每份作业的截止日期和扣分值，并按扣分从大到小排序。",
            "依次处理排序后的作业，优先保证高扣分作业按时完成。",
            "为当前作业寻找一个不晚于其截止日期的最靠后空闲时间槽。",
            "若找到可用时间槽，则安排该作业并更新空闲时间集合。",
            "若找不到可用时间槽，则将该作业的扣分加入总扣分。",
            "全部作业处理完成后，得到最少扣分值。"
        ],
        "complexity": (
            "排序需要 O(n log n) 时间；之后每份作业通过并查集查找和合并可用时间槽，"
            "均摊复杂度近似为常数，因此总时间复杂度为 O(n log n)。"
        ),
        "code_summary": (
            "核心程序位于 exp4_homework_scheduling/src/main.cpp。代码定义了作业结构体 Job，"
            "并用 DisjointSet 维护当前仍可使用的时间槽，实现“最晚可行安排”。"
        ),
        "test_summary": [
            "对指导书中的三组样例进行测试，程序输出 0、3、5，与标准输出完全一致。",
            "排序后优先安排高扣分作业，能够有效减少最终总罚分。",
            "并查集的加入使得每次安排时间槽的效率更高，适合较大规模输入。"
        ],
        "screenshots": [
            "截图 1：贪心排序和并查集安排逻辑的代码调试截图。",
            "截图 2：输入指导书三组样例后的运行结果截图。",
            "截图 3：程序输出最少扣分结果的截图。"
        ],
        "conclusion": (
            "通过本实验，掌握了调度问题中“局部最优选择”的设计方法，"
            "并理解了贪心算法适用的关键条件：贪心选择性质和最优子结构性质。"
        ),
    },
]


def replace_paragraph_text(paragraph, text):
    for run in paragraph.runs:
        run.text = ""
    if paragraph.runs:
        paragraph.runs[0].text = text
    else:
        paragraph.add_run(text)


def add_heading(doc, text, level):
    paragraph = doc.add_paragraph(text)
    paragraph.style = f"Heading {level}"
    return paragraph


def add_paragraph(doc, text="", align=None):
    paragraph = doc.add_paragraph(text)
    if align is not None:
        paragraph.alignment = align
    return paragraph


def add_list(doc, items):
    for idx, item in enumerate(items, start=1):
        add_paragraph(doc, f"{idx}. {item}")


def main():
    REPORT_DIR.mkdir(parents=True, exist_ok=True)

    for experiment in EXPERIMENTS:
        doc = Document(TEMPLATE)

        replace_paragraph_text(doc.paragraphs[14], "2026年 4 月")
        replace_paragraph_text(doc.paragraphs[15], experiment["title"])
        replace_paragraph_text(doc.paragraphs[16], experiment["description"])
        replace_paragraph_text(doc.paragraphs[17], "要求：")
        replace_paragraph_text(doc.paragraphs[18], "使用C++程序设计语言完成题目。")
        replace_paragraph_text(doc.paragraphs[19], "编写完程序后，回答如下问题：")
        replace_paragraph_text(doc.paragraphs[20], f"      {experiment['question1']}")
        replace_paragraph_text(doc.paragraphs[21], experiment["question2"])
        replace_paragraph_text(doc.paragraphs[23], "3、实验报告要求：直接在本文档中续写，要求粘贴代码调试截图和运行结果截图，回答相应问题。")
        replace_paragraph_text(doc.paragraphs[24], "")

        table = doc.tables[0]
        if not table.cell(4, 1).text.strip():
            table.cell(4, 1).text = "待填写"
        if not table.cell(5, 1).text.strip():
            table.cell(5, 1).text = "待填写"

        add_heading(doc, "四、实验设计与实现", 2)
        add_paragraph(doc, f"1. 算法策略：{experiment['strategy']}")
        add_paragraph(doc, f"2. 主要思想：{experiment['core']}")

        add_paragraph(doc, "3. 算法框架：")
        add_list(doc, experiment["framework"])

        add_paragraph(doc, f"4. 时间复杂度分析：{experiment['complexity']}")

        add_heading(doc, "五、程序实现说明", 2)
        add_paragraph(doc, experiment["code_summary"])

        add_heading(doc, "六、问题回答", 2)
        add_paragraph(doc, f"问题（1）：{experiment['question1']}")
        add_paragraph(doc, experiment["core"])
        add_paragraph(doc, "算法框架如下：")
        add_list(doc, experiment["framework"])
        add_paragraph(doc, f"问题（2）：{experiment['question2']}")
        add_paragraph(doc, experiment["complexity"])

        if "extra_question" in experiment:
            add_paragraph(doc, experiment["extra_question"]["title"])
            for line in experiment["extra_question"]["content"]:
                add_paragraph(doc, line)

        add_heading(doc, "七、测试结果与分析", 2)
        add_list(doc, experiment["test_summary"])

        add_heading(doc, "八、截图粘贴说明", 2)
        add_paragraph(doc, "按照指导书要求，建议至少补充以下截图：")
        add_list(doc, experiment["screenshots"])
        add_paragraph(doc, "【在此插入截图 1】", align=WD_PARAGRAPH_ALIGNMENT.CENTER)
        add_paragraph(doc, "【在此插入截图 2】", align=WD_PARAGRAPH_ALIGNMENT.CENTER)
        add_paragraph(doc, "【在此插入截图 3】", align=WD_PARAGRAPH_ALIGNMENT.CENTER)

        add_heading(doc, "九、实验总结", 2)
        add_paragraph(doc, experiment["conclusion"])

        output_path = REPORT_DIR / experiment["filename"]
        doc.save(output_path)
        print(f"Generated: {output_path}")


if __name__ == "__main__":
    main()

